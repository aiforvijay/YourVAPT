import streamlit as st
import pandas as pd
import requests
import json
import os
import shutil
from datetime import datetime
import base64
from io import BytesIO
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
from openai import OpenAI

# Page configuration
st.set_page_config(
    page_title="Your System VAPT Analysis",
    page_icon="🔒",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Initialize session state
if 'api_key' not in st.session_state:
    st.session_state.api_key = ""
if 'command_outputs' not in st.session_state:
    st.session_state.command_outputs = {}
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = {}
if 'progress' not in st.session_state:
    st.session_state.progress = {}

# Create data directory if it doesn't exist
os.makedirs("data", exist_ok=True)

# Windows 11 Commands Database
WINDOWS_COMMANDS = {
    "💻 System Information": [
        {
            "command": "systeminfo",
            "purpose": "Reveals OS version, build, hotfixes, and system configuration (risk: unpatched vulnerabilities)",
            "instructions": "Open Command Prompt as Administrator → Type 'systeminfo' → Press Enter → Wait for system information to display",
            "shortcuts": "Win+X → Command Prompt (Admin) or Win+R → cmd → Ctrl+Shift+Enter"
        },
        {
            "command": "ver",
            "purpose": "Shows Windows version number (risk: outdated OS versions)",
            "instructions": "Open Command Prompt → Type 'ver' → Press Enter",
            "shortcuts": "Win+R → cmd → Enter"
        },
        {
            "command": "Get-ComputerInfo",
            "purpose": "PowerShell command for detailed system information (risk: configuration weaknesses)",
            "instructions": "Open PowerShell as Administrator → Type 'Get-ComputerInfo' → Press Enter",
            "shortcuts": "Win+X → Windows PowerShell (Admin)"
        },
        {
            "command": "wmic os get Caption,Version,BuildNumber",
            "purpose": "Gets OS details via WMI (risk: version-specific vulnerabilities)",
            "instructions": "Open Command Prompt → Type the command → Press Enter",
            "shortcuts": "Win+R → cmd → Enter"
        }
    ],
    "👥 User Management": [
        {
            "command": "net user",
            "purpose": "Lists all user accounts on the system (risk: unauthorized accounts, weak passwords)",
            "instructions": "Open Command Prompt as Administrator → Type 'net user' → Press Enter",
            "shortcuts": "Win+X → Command Prompt (Admin)"
        },
        {
            "command": "net localgroup administrators",
            "purpose": "Shows administrator group members (risk: privilege escalation, unauthorized admin access)",
            "instructions": "Open Command Prompt as Administrator → Type 'net localgroup administrators' → Press Enter",
            "shortcuts": "Win+X → Command Prompt (Admin)"
        },
        {
            "command": "Get-LocalUser | Select Name,Enabled,LastLogon",
            "purpose": "PowerShell command to list local users with status (risk: inactive accounts, service accounts)",
            "instructions": "Open PowerShell as Administrator → Type the command → Press Enter",
            "shortcuts": "Win+X → Windows PowerShell (Admin)"
        },
        {
            "command": "whoami /priv",
            "purpose": "Shows current user privileges (risk: excessive privileges, privilege abuse)",
            "instructions": "Open Command Prompt → Type 'whoami /priv' → Press Enter",
            "shortcuts": "Win+R → cmd → Enter"
        }
    ],
    "🔒 Network Security": [
        {
            "command": "netstat -an",
            "purpose": "Shows all active network connections and listening ports (risk: unauthorized services, backdoors)",
            "instructions": "Open Command Prompt → Type 'netstat -an' → Press Enter",
            "shortcuts": "Win+R → cmd → Enter"
        },
        {
            "command": "netsh wlan show profiles",
            "purpose": "Displays saved WiFi profiles (risk: weak encryption, rogue networks)",
            "instructions": "Open Command Prompt → Type 'netsh wlan show profiles' → Press Enter",
            "shortcuts": "Win+R → cmd → Enter"
        },
        {
            "command": "ipconfig /all",
            "purpose": "Shows network configuration details (risk: DNS poisoning, network misconfigurations)",
            "instructions": "Open Command Prompt → Type 'ipconfig /all' → Press Enter",
            "shortcuts": "Win+R → cmd → Enter"
        },
        {
            "command": "arp -a",
            "purpose": "Displays ARP table entries (risk: ARP spoofing, man-in-the-middle attacks)",
            "instructions": "Open Command Prompt → Type 'arp -a' → Press Enter",
            "shortcuts": "Win+R → cmd → Enter"
        }
    ],
    "🛡️ Security Software": [
        {
            "command": "Get-MpPreference",
            "purpose": "Shows Windows Defender settings (risk: disabled protection, exclusions)",
            "instructions": "Open PowerShell as Administrator → Type 'Get-MpPreference' → Press Enter",
            "shortcuts": "Win+X → Windows PowerShell (Admin)"
        },
        {
            "command": "Get-MpComputerStatus",
            "purpose": "Shows Windows Defender status (risk: outdated definitions, disabled real-time protection)",
            "instructions": "Open PowerShell as Administrator → Type 'Get-MpComputerStatus' → Press Enter",
            "shortcuts": "Win+X → Windows PowerShell (Admin)"
        },
        {
            "command": "wmic /namespace:\\\\root\\SecurityCenter2 path AntiVirusProduct get *",
            "purpose": "Lists installed antivirus products (risk: no antivirus, multiple conflicting AV)",
            "instructions": "Open Command Prompt → Type the command → Press Enter",
            "shortcuts": "Win+R → cmd → Enter"
        },
        {
            "command": "sc query windefend",
            "purpose": "Checks Windows Defender service status (risk: disabled security services)",
            "instructions": "Open Command Prompt → Type 'sc query windefend' → Press Enter",
            "shortcuts": "Win+R → cmd → Enter"
        }
    ],
    "⚠️ Application Security": [
        {
            "command": "wmic product get name,version",
            "purpose": "Lists installed software with versions (risk: vulnerable applications, unlicensed software)",
            "instructions": "Open Command Prompt → Type 'wmic product get name,version' → Press Enter (may take time)",
            "shortcuts": "Win+R → cmd → Enter"
        },
        {
            "command": "Get-AppxPackage | Select Name,Version",
            "purpose": "Lists installed Windows Store apps (risk: vulnerable apps, unnecessary permissions)",
            "instructions": "Open PowerShell → Type the command → Press Enter",
            "shortcuts": "Win+X → Windows PowerShell"
        },
        {
            "command": "sfc /verifyonly",
            "purpose": "Verifies system file integrity (risk: corrupted system files, malware)",
            "instructions": "Open Command Prompt as Administrator → Type 'sfc /verifyonly' → Press Enter",
            "shortcuts": "Win+X → Command Prompt (Admin)"
        },
        {
            "command": "powershell Get-ExecutionPolicy",
            "purpose": "Shows PowerShell execution policy (risk: unrestricted script execution)",
            "instructions": "Open Command Prompt → Type 'powershell Get-ExecutionPolicy' → Press Enter",
            "shortcuts": "Win+R → cmd → Enter"
        }
    ],
    "💾 Data Protection": [
        {
            "command": "manage-bde -status",
            "purpose": "Shows BitLocker encryption status (risk: unencrypted drives, weak encryption)",
            "instructions": "Open Command Prompt as Administrator → Type 'manage-bde -status' → Press Enter",
            "shortcuts": "Win+X → Command Prompt (Admin)"
        },
        {
            "command": "vssadmin list shadows",
            "purpose": "Lists volume shadow copies (risk: data recovery attacks, backup vulnerabilities)",
            "instructions": "Open Command Prompt as Administrator → Type 'vssadmin list shadows' → Press Enter",
            "shortcuts": "Win+X → Command Prompt (Admin)"
        },
        {
            "command": "fsutil behavior query DisableDeleteNotify",
            "purpose": "Checks TRIM support for SSDs (risk: data recovery from deleted files)",
            "instructions": "Open Command Prompt as Administrator → Type the command → Press Enter",
            "shortcuts": "Win+X → Command Prompt (Admin)"
        },
        {
            "command": "cipher /c",
            "purpose": "Shows encryption status of files and folders (risk: unencrypted sensitive data)",
            "instructions": "Open Command Prompt as Administrator → Type 'cipher /c' → Press Enter",
            "shortcuts": "Win+X → Command Prompt (Admin)"
        }
    ],
    "📋 Logging and Monitoring": [
        {
            "command": "auditpol /get /category:*",
            "purpose": "Shows audit policy settings (risk: insufficient logging, compliance violations)",
            "instructions": "Open Command Prompt as Administrator → Type 'auditpol /get /category:*' → Press Enter",
            "shortcuts": "Win+X → Command Prompt (Admin)"
        },
        {
            "command": "wevtutil el",
            "purpose": "Lists available event logs (risk: missing security logs, log tampering)",
            "instructions": "Open Command Prompt → Type 'wevtutil el' → Press Enter",
            "shortcuts": "Win+R → cmd → Enter"
        },
        {
            "command": "Get-EventLog -List",
            "purpose": "PowerShell command to list event logs (risk: log analysis blind spots)",
            "instructions": "Open PowerShell → Type 'Get-EventLog -List' → Press Enter",
            "shortcuts": "Win+X → Windows PowerShell"
        },
        {
            "command": "wmic nteventlog get LogFileName,MaxFileSize",
            "purpose": "Shows event log file settings (risk: log size limits, log rotation issues)",
            "instructions": "Open Command Prompt → Type the command → Press Enter",
            "shortcuts": "Win+R → cmd → Enter"
        }
    ],
    "🦠 Antivirus & Threat Detection": [
        {
            "command": "$data = Get-MpThreatDetection; if ($data) { $data | Export-Csv -Path \"C:\\DefenderThreats.csv\" -NoTypeInformation; Write-Host \"Exported $($data.Count) detection(s) to C:\\DefenderThreats.csv\" } else { Write-Warning \"No Defender threat detections found to export.\" }",
            "purpose": "Exports Windows Defender threat detections or confirms no threats found (risk: active malware, false negatives)",
            "instructions": "Open PowerShell as Administrator → Copy and paste the full command → Press Enter",
            "shortcuts": "Win+X → Windows PowerShell (Admin)"
        },
        {
            "command": "Get-MpComputerStatus",
            "purpose": "Shows Windows Defender status, definitions, last scan time (risk: outdated definitions, disabled protection)",
            "instructions": "Open PowerShell as Administrator → Type 'Get-MpComputerStatus' → Press Enter",
            "shortcuts": "Win+X → Windows PowerShell (Admin)"
        },
        {
            "command": "schtasks /query /fo LIST | findstr /i defender",
            "purpose": "Shows Windows Defender scheduled tasks (risk: disabled automatic scans)",
            "instructions": "Open Command Prompt → Type the command → Press Enter",
            "shortcuts": "Win+R → cmd → Enter"
        },
        {
            "command": "reg query \"HKLM\\SOFTWARE\\Microsoft\\Windows Defender\" /s",
            "purpose": "Shows Windows Defender registry settings (risk: tampered security settings)",
            "instructions": "Open Command Prompt as Administrator → Type the command → Press Enter",
            "shortcuts": "Win+X → Command Prompt (Admin)"
        },
        {
            "command": "Get-NetFirewallProfile | Format-Table Name, Enabled, DefaultInboundAction, DefaultOutboundAction",
            "purpose": "Shows firewall profile information and default actions (risk: disabled firewall profiles, permissive default actions)",
            "instructions": "Open PowerShell as Administrator → Type the command → Press Enter",
            "shortcuts": "Win+X → Windows PowerShell (Admin)"
        }
    ],
    "🔥 Firewall Configuration": [
        {
            "command": "netsh advfirewall show allprofiles",
            "purpose": "Shows Windows Firewall status for all profiles (risk: disabled firewall, weak rules)",
            "instructions": "Open Command Prompt as Administrator → Type the command → Press Enter",
            "shortcuts": "Win+X → Command Prompt (Admin)"
        },
        {
            "command": "netsh advfirewall firewall show rule name=all",
            "purpose": "Lists all firewall rules (risk: overly permissive rules, backdoor rules)",
            "instructions": "Open Command Prompt as Administrator → Type the command → Press Enter",
            "shortcuts": "Win+X → Command Prompt (Admin)"
        },
        {
            "command": "Get-NetFirewallProfile",
            "purpose": "PowerShell command for firewall profile status (risk: inconsistent profile settings)",
            "instructions": "Open PowerShell as Administrator → Type 'Get-NetFirewallProfile' → Press Enter",
            "shortcuts": "Win+X → Windows PowerShell (Admin)"
        },
        {
            "command": "netsh firewall show config",
            "purpose": "Shows legacy firewall configuration (risk: legacy firewall conflicts)",
            "instructions": "Open Command Prompt as Administrator → Type 'netsh firewall show config' → Press Enter",
            "shortcuts": "Win+X → Command Prompt (Admin)"
        }
    ]
}

# macOS Commands Database
MACOS_COMMANDS = {
    "💻 System Information": [
        {
            "command": "sw_vers",
            "purpose": "Shows macOS version and build information (risk: unpatched vulnerabilities)",
            "instructions": "Open Terminal → Type 'sw_vers' → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "system_profiler SPSoftwareDataType",
            "purpose": "Detailed system software information (risk: outdated system components)",
            "instructions": "Open Terminal → Type 'system_profiler SPSoftwareDataType' → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "uname -a",
            "purpose": "Shows kernel and system information (risk: kernel vulnerabilities)",
            "instructions": "Open Terminal → Type 'uname -a' → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "sysctl -a | grep machdep.cpu",
            "purpose": "Shows CPU information and features (risk: CPU vulnerabilities, missing mitigations)",
            "instructions": "Open Terminal → Type the command → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        }
    ],
    "👥 User Management": [
        {
            "command": "dscl . -list /Users",
            "purpose": "Lists all user accounts on the system (risk: unauthorized accounts)",
            "instructions": "Open Terminal → Type 'dscl . -list /Users' → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "dscl . -read /Groups/admin GroupMembership",
            "purpose": "Shows administrator group members (risk: excessive admin privileges)",
            "instructions": "Open Terminal → Type the command → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "id",
            "purpose": "Shows current user ID and group memberships (risk: privilege verification)",
            "instructions": "Open Terminal → Type 'id' → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "last | head -20",
            "purpose": "Shows recent user login history (risk: unauthorized access, suspicious logins)",
            "instructions": "Open Terminal → Type 'last | head -20' → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        }
    ],
    "🔒 Network Security": [
        {
            "command": "netstat -an",
            "purpose": "Shows all active network connections and listening ports (risk: unauthorized services)",
            "instructions": "Open Terminal → Type 'netstat -an' → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "security find-generic-password -wa",
            "purpose": "Lists WiFi passwords from keychain (risk: weak WiFi security)",
            "instructions": "Open Terminal → Type the command → Enter admin password when prompted",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "ifconfig",
            "purpose": "Shows network interface configurations (risk: network misconfigurations)",
            "instructions": "Open Terminal → Type 'ifconfig' → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "arp -a",
            "purpose": "Displays ARP table entries (risk: ARP spoofing attacks)",
            "instructions": "Open Terminal → Type 'arp -a' → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        }
    ],
    "🛡️ Security Software": [
        {
            "command": "system_profiler SPApplicationsDataType | grep -i antivirus",
            "purpose": "Checks for installed antivirus software (risk: no antivirus protection)",
            "instructions": "Open Terminal → Type the command → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "spctl --status",
            "purpose": "Shows Gatekeeper status (risk: disabled app security)",
            "instructions": "Open Terminal → Type 'spctl --status' → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "csrutil status",
            "purpose": "Shows System Integrity Protection status (risk: disabled SIP)",
            "instructions": "Open Terminal → Type 'csrutil status' → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "kextstat | grep -v apple",
            "purpose": "Lists non-Apple kernel extensions (risk: malicious kernel extensions)",
            "instructions": "Open Terminal → Type 'kextstat | grep -v apple' → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        }
    ],
    "⚠️ Application Security": [
        {
            "command": "system_profiler SPApplicationsDataType",
            "purpose": "Lists all installed applications (risk: vulnerable applications)",
            "instructions": "Open Terminal → Type 'system_profiler SPApplicationsDataType' → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "launchctl list | grep -v com.apple",
            "purpose": "Shows non-Apple launch agents and daemons (risk: malicious processes)",
            "instructions": "Open Terminal → Type the command → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "defaults read com.apple.security.libraryvalidation.plist",
            "purpose": "Checks library validation settings (risk: code injection attacks)",
            "instructions": "Open Terminal → Type the command → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "codesign --verify --deep --strict /Applications/*.app",
            "purpose": "Verifies application signatures (risk: tampered applications)",
            "instructions": "Open Terminal → Type the command → Press Enter (may take time)",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        }
    ],
    "💾 Data Protection": [
        {
            "command": "fdesetup status",
            "purpose": "Shows FileVault encryption status (risk: unencrypted drives)",
            "instructions": "Open Terminal → Type 'fdesetup status' → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "tmutil status",
            "purpose": "Shows Time Machine backup status (risk: no backups, failed backups)",
            "instructions": "Open Terminal → Type 'tmutil status' → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "security list-keychains",
            "purpose": "Lists available keychains (risk: keychain vulnerabilities)",
            "instructions": "Open Terminal → Type 'security list-keychains' → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "diskutil list",
            "purpose": "Shows disk partitions and encryption status (risk: unencrypted partitions)",
            "instructions": "Open Terminal → Type 'diskutil list' → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        }
    ],
    "📋 Logging and Monitoring": [
        {
            "command": "sudo log show --last 1h --predicate 'eventType == logEvent'",
            "purpose": "Shows recent system logs (risk: insufficient logging)",
            "instructions": "Open Terminal → Type the command → Enter admin password",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "ls -la /var/log/",
            "purpose": "Lists available log files (risk: missing security logs)",
            "instructions": "Open Terminal → Type 'ls -la /var/log/' → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "sudo fs_usage -w | head -50",
            "purpose": "Shows file system usage in real-time (risk: unauthorized file access)",
            "instructions": "Open Terminal → Type the command → Enter admin password → Ctrl+C to stop",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "pmset -g log | grep -i wake",
            "purpose": "Shows power management logs (risk: unexpected system wake events)",
            "instructions": "Open Terminal → Type 'pmset -g log | grep -i wake' → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        }
    ],
    "🦠 Antivirus & Threat Detection": [
        {
            "command": "sudo malware-scanner --scan /",
            "purpose": "Runs system-wide malware scan (risk: undetected malware)",
            "instructions": "Open Terminal → Type the command → Enter admin password (if scanner installed)",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "system_profiler SPInstallHistoryDataType | grep -i security",
            "purpose": "Shows security update history (risk: missing security updates)",
            "instructions": "Open Terminal → Type the command → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "sudo lsof -i | grep LISTEN",
            "purpose": "Shows processes listening on network ports (risk: malicious network services)",
            "instructions": "Open Terminal → Type the command → Enter admin password",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "ps aux | grep -v '] 0'",
            "purpose": "Shows running processes (risk: malicious processes)",
            "instructions": "Open Terminal → Type 'ps aux | grep -v \"] 0\"' → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        }
    ],
    "🔥 Firewall Configuration": [
        {
            "command": "sudo /usr/libexec/ApplicationFirewall/socketfilterfw --getglobalstate",
            "purpose": "Shows macOS firewall status (risk: disabled firewall)",
            "instructions": "Open Terminal → Type the command → Enter admin password",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "sudo /usr/libexec/ApplicationFirewall/socketfilterfw --listapps",
            "purpose": "Lists firewall application rules (risk: unauthorized network access)",
            "instructions": "Open Terminal → Type the command → Enter admin password",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "sudo pfctl -s rules",
            "purpose": "Shows packet filter rules (risk: weak firewall rules)",
            "instructions": "Open Terminal → Type 'sudo pfctl -s rules' → Enter admin password",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        },
        {
            "command": "netstat -rn",
            "purpose": "Shows routing table (risk: routing vulnerabilities)",
            "instructions": "Open Terminal → Type 'netstat -rn' → Press Enter",
            "shortcuts": "Cmd+Space → Type 'Terminal' → Enter"
        }
    ]
}

def detect_ai_provider(api_key):
    """Detect AI provider based on API key format"""
    if not api_key:
        return None
    
    if api_key.startswith('sk-'):
        return 'openai'
    elif api_key.startswith('gsk_'):
        return 'openai'  # Some OpenAI keys start with gsk_
    elif 'claude' in api_key.lower() or api_key.startswith('sk-ant'):
        return 'anthropic'
    elif 'AIza' in api_key:
        return 'gemini'
    elif 'xai-' in api_key.lower():
        return 'grok'
    else:
        return 'openai'  # Default to OpenAI format

def clean_ai_response(response_text):
    """Clean up AI response by removing preambles and markdown formatting"""
    if not response_text:
        return response_text
    
    # Remove common AI preambles
    preambles_to_remove = [
        "Of course. Here is a detailed cybersecurity analysis of the provided system command output.",
        "Here is a detailed cybersecurity analysis of the provided system command output:",
        "I'll analyze this system command output for cybersecurity vulnerabilities and risks:",
        "Based on the provided command output, here is my analysis:",
        "Here's my analysis of the command output:",
        "Certainly! Here is a detailed analysis:",
        "Of course! Here's the analysis:",
        "Here is the vulnerability analysis:"
    ]
    
    cleaned_text = response_text.strip()
    
    # Remove preambles (case insensitive)
    for preamble in preambles_to_remove:
        if cleaned_text.lower().startswith(preamble.lower()):
            cleaned_text = cleaned_text[len(preamble):].strip()
            break
    
    # Remove markdown formatting
    import re
    # Remove markdown headers (##, ###, etc.)
    cleaned_text = re.sub(r'^#+\s*', '', cleaned_text, flags=re.MULTILINE)
    # Remove bold formatting (**text**)
    cleaned_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cleaned_text)
    # Remove italic formatting (*text*)
    cleaned_text = re.sub(r'\*(.*?)\*', r'\1', cleaned_text)
    # Remove markdown links [text](url)
    cleaned_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', cleaned_text)
    # Remove backticks
    cleaned_text = re.sub(r'`([^`]+)`', r'\1', cleaned_text)
    # Remove triple backticks and code blocks
    cleaned_text = re.sub(r'```[\s\S]*?```', '', cleaned_text)
    
    # Clean up extra whitespace
    cleaned_text = re.sub(r'\n\s*\n\s*\n', '\n\n', cleaned_text)
    cleaned_text = cleaned_text.strip()
    
    # Ensure it starts with "Executive Summary" if it doesn't already
    if not cleaned_text.lower().startswith('executive summary'):
        if cleaned_text:
            cleaned_text = "Executive Summary\n\n" + cleaned_text
    
    return cleaned_text

def analyze_with_ai(output_text, api_key, section_name):
    """Analyze command output using AI"""
    try:
        provider = detect_ai_provider(api_key)
        
        prompt = f"""
        Analyze this system command output for cybersecurity vulnerabilities and risks:
        
        Section: {section_name}
        Command Output: {output_text}
        
        Please provide a professional vulnerability assessment report starting with "Executive Summary" and including:
        1. Risk Score (Low/Medium/High)
        2. CVSS Score (0-10) if applicable
        3. Specific vulnerabilities found
        4. CVE references if applicable
        5. References to security frameworks (NIST SP 800-53, ISO 27001, OWASP Top 10)
        6. Recommended remediation steps
        
        Format your response as clean, professional text without markdown formatting or preambles.
        """
        
        if provider == 'openai':
            client = OpenAI(api_key=api_key)
            # the newest OpenAI model is "gpt-5" which was released August 7, 2025.
            # do not change this unless explicitly requested by the user
            response = client.chat.completions.create(
                model="gpt-5",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1000
            )
            raw_response = response.choices[0].message.content
            return clean_ai_response(raw_response)
            
        elif provider == 'anthropic':
            # Using requests to call Anthropic API
            headers = {
                "x-api-key": api_key,
                "Content-Type": "application/json",
                "anthropic-version": "2023-06-01"
            }
            data = {
                "model": "claude-sonnet-4-20250514",
                "max_tokens": 1000,
                "messages": [{"role": "user", "content": prompt}]
            }
            response = requests.post(
                "https://api.anthropic.com/v1/messages",
                headers=headers,
                json=data
            )
            if response.status_code == 200:
                raw_response = response.json()["content"][0]["text"]
                return clean_ai_response(raw_response)
            else:
                return f"Error: {response.status_code} - {response.text}"
                
        elif provider == 'gemini':
            # Using requests to call Gemini API
            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-pro:generateContent?key={api_key}"
            data = {
                "contents": [{"parts": [{"text": prompt}]}]
            }
            response = requests.post(url, json=data)
            if response.status_code == 200:
                raw_response = response.json()["candidates"][0]["content"]["parts"][0]["text"]
                return clean_ai_response(raw_response)
            else:
                return f"Error: {response.status_code} - {response.text}"
                
        elif provider == 'grok':
            # Using OpenAI-compatible API for Grok
            client = OpenAI(
                api_key=api_key,
                base_url="https://api.x.ai/v1"
            )
            response = client.chat.completions.create(
                model="grok-2-1212",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1000
            )
            raw_response = response.choices[0].message.content
            return clean_ai_response(raw_response)
            
    except Exception as e:
        return f"AI Analysis Error: {str(e)}"

def mock_analyze_output(output_text, section_name):
    """Provide mock analysis for demonstration"""
    mock_analyses = {
        "💻 System Information": {
            "risk": "Medium",
            "cvss": "6.5",
            "vulnerabilities": "Outdated OS version detected, missing security patches",
            "cve": "CVE-2023-1234, CVE-2023-5678",
            "references": "NIST SP 800-53 SI-2, ISO 27001 A.12.6.1, OWASP Top 10 A06",
            "remediation": "Apply latest security updates, enable automatic updates"
        },
        "👥 User Management": {
            "risk": "High", 
            "cvss": "8.2",
            "vulnerabilities": "Multiple administrator accounts detected, guest account enabled",
            "cve": "N/A",
            "references": "NIST SP 800-53 AC-2, ISO 27001 A.9.2.1, OWASP Top 10 A07",
            "remediation": "Disable unnecessary admin accounts, implement principle of least privilege"
        },
        "🔒 Network Security": {
            "risk": "Medium",
            "cvss": "5.8",
            "vulnerabilities": "Open ports detected, weak WiFi security protocols",
            "cve": "N/A",
            "references": "NIST SP 800-53 SC-7, ISO 27001 A.13.1.1, OWASP Top 10 A05",
            "remediation": "Close unnecessary ports, upgrade to WPA3 encryption"
        },
        "🛡️ Security Software": {
            "risk": "High",
            "cvss": "7.8",
            "vulnerabilities": "Antivirus disabled, real-time protection off",
            "cve": "N/A", 
            "references": "NIST SP 800-53 SI-3, ISO 27001 A.12.2.1, OWASP Top 10 A06",
            "remediation": "Enable antivirus, update virus definitions, enable real-time scanning"
        },
        "⚠️ Application Security": {
            "risk": "Medium",
            "cvss": "6.1",
            "vulnerabilities": "Vulnerable software versions detected, unsigned applications",
            "cve": "CVE-2023-9999",
            "references": "NIST SP 800-53 SI-2, ISO 27001 A.12.6.1, OWASP Top 10 A06",
            "remediation": "Update vulnerable applications, implement application whitelisting"
        },
        "💾 Data Protection": {
            "risk": "High",
            "cvss": "8.5",
            "vulnerabilities": "Drive encryption disabled, no backup strategy",
            "cve": "N/A",
            "references": "NIST SP 800-53 SC-28, ISO 27001 A.10.1.1, OWASP Top 10 A02",
            "remediation": "Enable full disk encryption, implement automated backups"
        },
        "📋 Logging and Monitoring": {
            "risk": "Medium",
            "cvss": "5.2",
            "vulnerabilities": "Insufficient audit logging, log retention too short",
            "cve": "N/A",
            "references": "NIST SP 800-53 AU-2, ISO 27001 A.12.4.1, OWASP Top 10 A09",
            "remediation": "Enable comprehensive audit logging, extend log retention period"
        },
        "🦠 Antivirus & Threat Detection": {
            "risk": "High",
            "cvss": "8.0",
            "vulnerabilities": "No recent malware scans, threat detection disabled",
            "cve": "N/A",
            "references": "NIST SP 800-53 SI-3, ISO 27001 A.12.2.1, OWASP Top 10 A06",
            "remediation": "Run full system scan, enable automatic threat detection"
        },
        "🔥 Firewall Configuration": {
            "risk": "High",
            "cvss": "7.5",
            "vulnerabilities": "Firewall disabled, overly permissive rules",
            "cve": "N/A",
            "references": "NIST SP 800-53 SC-7, ISO 27001 A.13.1.1, OWASP Top 10 A05",
            "remediation": "Enable firewall, implement restrictive default-deny rules"
        }
    }
    
    return mock_analyses.get(section_name, {
        "risk": "Low",
        "cvss": "3.0", 
        "vulnerabilities": "No significant vulnerabilities detected in sample analysis",
        "cve": "N/A",
        "references": "NIST SP 800-53, ISO 27001, OWASP Top 10",
        "remediation": "Continue monitoring and follow security best practices"
    })

def calculate_overall_progress():
    """Calculate overall progress based on completed sections"""
    if not st.session_state.progress:
        return 0
    
    total_sections = 18  # 9 sections per OS
    completed = sum(1 for status in st.session_state.progress.values() if status)
    return int((completed / total_sections) * 100)

def generate_pdf_report(auditor_name, audit_cert, organization, assessment_date, selected_os):
    """Generate PDF report with comprehensive vulnerability analysis"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=0.75*inch, rightMargin=0.75*inch, 
                           topMargin=1*inch, bottomMargin=1*inch)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        spaceAfter=20,
        alignment=1,  # Center alignment
        textColor=colors.darkblue,
        fontName='Helvetica-Bold'
    )
    story.append(Paragraph("VAPT Analysis Report", title_style))
    story.append(Spacer(1, 20))
    
    # Report Details Table - styled like the sample
    details = [
        ["Auditor Name", auditor_name or "Not Specified"],
        ["Certificate #", audit_cert or "Not Specified"],
        ["Organization", organization or "Not Specified"], 
        ["Assessment Date", str(assessment_date)],
        ["Target OS", selected_os],
        ["Report Prepared", datetime.now().strftime("%d/%m/%Y")]
    ]
    
    details_table = Table(details, colWidths=[2.5*inch, 3.5*inch])
    details_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightblue),
        ('BACKGROUND', (1, 0), (1, -1), colors.white),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
    ]))
    
    story.append(details_table)
    story.append(Spacer(1, 30))
    
    # Analysis Results Section
    analysis_header_style = ParagraphStyle(
        'AnalysisHeader',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=15,
        textColor=colors.darkblue,
        fontName='Helvetica-Bold'
    )
    story.append(Paragraph("Analysis Results", analysis_header_style))
    
    # Check if we have any analysis results
    if st.session_state.analysis_results:
        section_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading3'],
            fontSize=12,
            spaceAfter=8,
            spaceBefore=15,
            textColor=colors.darkgreen,
            fontName='Helvetica-Bold'
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=4,
            leftIndent=20
        )
        
        # Process each section's analysis results
        for section_key, result in st.session_state.analysis_results.items():
            # Extract clean section name (remove os prefix)
            section_name = section_key.replace('windows_', '').replace('macos_', '').replace('_', ' ').title()
            story.append(Paragraph(f"{section_name}", section_style))
            
            if isinstance(result, dict):
                # Structured result from Test Run
                story.append(Paragraph(f"<b>Risk Level:</b> {result.get('risk', 'N/A')}", normal_style))
                story.append(Paragraph(f"<b>CVSS Score:</b> {result.get('cvss', 'N/A')}", normal_style))
                story.append(Paragraph(f"<b>Vulnerabilities Found:</b> {result.get('vulnerabilities', 'N/A')}", normal_style))
                
                if result.get('cve', 'N/A') != 'N/A':
                    story.append(Paragraph(f"<b>CVE References:</b> {result.get('cve')}", normal_style))
                
                story.append(Paragraph(f"<b>Framework References:</b> {result.get('references', 'N/A')}", normal_style))
                story.append(Paragraph(f"<b>Recommended Actions:</b> {result.get('remediation', 'N/A')}", normal_style))
                
            else:
                # Text result from AI Analysis - format it properly
                result_text = str(result).strip()
                if result_text:
                    # Split long text into paragraphs for better formatting
                    paragraphs = result_text.split('\n')
                    for para in paragraphs:
                        if para.strip():
                            story.append(Paragraph(para.strip(), normal_style))
                else:
                    story.append(Paragraph("Analysis completed - no specific vulnerabilities identified.", normal_style))
            
            story.append(Spacer(1, 10))
    
    else:
        # No analysis results found
        no_results_style = ParagraphStyle(
            'NoResults',
            parent=styles['Normal'],
            fontSize=11,
            textColor=colors.red,
            alignment=1
        )
        story.append(Paragraph("No vulnerability analysis results found. Please run analysis on command outputs before generating the report.", no_results_style))
    
    # Summary section if we have results
    if st.session_state.analysis_results:
        story.append(Spacer(1, 20))
        summary_style = ParagraphStyle(
            'Summary',
            parent=styles['Heading3'],
            fontSize=12,
            spaceAfter=10,
            textColor=colors.darkblue,
            fontName='Helvetica-Bold'
        )
        story.append(Paragraph("Assessment Summary", summary_style))
        
        total_sections = len(st.session_state.analysis_results)
        high_risk = sum(1 for r in st.session_state.analysis_results.values() 
                       if isinstance(r, dict) and r.get('risk') == 'High')
        medium_risk = sum(1 for r in st.session_state.analysis_results.values() 
                         if isinstance(r, dict) and r.get('risk') == 'Medium')
        low_risk = sum(1 for r in st.session_state.analysis_results.values() 
                      if isinstance(r, dict) and r.get('risk') == 'Low')
        
        summary_text = f"Total sections analyzed: {total_sections}<br/>"
        if high_risk > 0:
            summary_text += f"High risk findings: {high_risk}<br/>"
        if medium_risk > 0:
            summary_text += f"Medium risk findings: {medium_risk}<br/>"
        if low_risk > 0:
            summary_text += f"Low risk findings: {low_risk}<br/>"
        
        summary_text += f"Overall assessment progress: {calculate_overall_progress()}%"
        
        story.append(Paragraph(summary_text, normal_style))
    
    # Footer
    story.append(Spacer(1, 30))
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.grey,
        alignment=1
    )
    story.append(Paragraph(f"Report generated on {datetime.now().strftime('%d/%m/%Y at %H:%M')} | VAPT Analysis Tool", footer_style))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def generate_docx_report(auditor_name, audit_cert, organization, assessment_date, selected_os):
    """Generate Word document report with comprehensive vulnerability analysis"""
    doc = Document()
    
    # Title
    title = doc.add_heading('VAPT Analysis Report', 0)
    title.alignment = 1  # Center alignment
    
    # Report details
    doc.add_heading('Report Details', level=1)
    details_table = doc.add_table(rows=5, cols=2)
    details_table.style = 'Table Grid'
    
    details_data = [
        ("Auditor Name", auditor_name or "Not Specified"),
        ("Organization", organization or "Not Specified"),
        ("Assessment Date", str(assessment_date)),
        ("Target OS", selected_os),
        ("Report Prepared", datetime.now().strftime("%d/%m/%Y"))
    ]
    
    for i, (label, value) in enumerate(details_data):
        details_table.cell(i, 0).text = label
        details_table.cell(i, 1).text = str(value)
        # Make first column bold
        details_table.cell(i, 0).paragraphs[0].runs[0].bold = True
    
    # Analysis Results
    doc.add_heading('Analysis Results', level=1)
    
    if st.session_state.analysis_results:
        for section_key, result in st.session_state.analysis_results.items():
            # Extract clean section name (remove os prefix)
            section_name = section_key.replace('windows_', '').replace('macos_', '').replace('_', ' ').title()
            doc.add_heading(section_name, level=2)
            
            if isinstance(result, dict):
                # Structured result from Test Run
                doc.add_paragraph(f"Risk Level: {result.get('risk', 'N/A')}")
                doc.add_paragraph(f"CVSS Score: {result.get('cvss', 'N/A')}")
                doc.add_paragraph(f"Vulnerabilities Found: {result.get('vulnerabilities', 'N/A')}")
                if result.get('cve', 'N/A') != 'N/A':
                    doc.add_paragraph(f"CVE References: {result.get('cve')}")
                doc.add_paragraph(f"Framework References: {result.get('references', 'N/A')}")
                doc.add_paragraph(f"Recommended Actions: {result.get('remediation', 'N/A')}")
            else:
                # Text result from AI Analysis
                result_text = str(result).strip()
                if result_text:
                    # Split long text into paragraphs
                    paragraphs = result_text.split('\n')
                    for para in paragraphs:
                        if para.strip():
                            doc.add_paragraph(para.strip())
                else:
                    doc.add_paragraph("Analysis completed - no specific vulnerabilities identified.")
    else:
        doc.add_paragraph("No vulnerability analysis results found. Please run analysis on command outputs before generating the report.")
    
    # Summary section
    if st.session_state.analysis_results:
        doc.add_heading('Assessment Summary', level=1)
        
        total_sections = len(st.session_state.analysis_results)
        high_risk = sum(1 for r in st.session_state.analysis_results.values() 
                       if isinstance(r, dict) and r.get('risk') == 'High')
        medium_risk = sum(1 for r in st.session_state.analysis_results.values() 
                         if isinstance(r, dict) and r.get('risk') == 'Medium')
        low_risk = sum(1 for r in st.session_state.analysis_results.values() 
                      if isinstance(r, dict) and r.get('risk') == 'Low')
        
        doc.add_paragraph(f"Total sections analyzed: {total_sections}")
        if high_risk > 0:
            doc.add_paragraph(f"High risk findings: {high_risk}")
        if medium_risk > 0:
            doc.add_paragraph(f"Medium risk findings: {medium_risk}")
        if low_risk > 0:
            doc.add_paragraph(f"Low risk findings: {low_risk}")
        doc.add_paragraph(f"Overall assessment progress: {calculate_overall_progress()}%")
    
    # Footer
    footer_para = doc.add_paragraph(f"Report generated on {datetime.now().strftime('%d/%m/%Y at %H:%M')} | VAPT Analysis Tool")
    footer_para.alignment = 1  # Center alignment
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_text_report(auditor_name, audit_cert, organization, assessment_date, selected_os):
    """Generate text report"""
    report_lines = []
    report_lines.append("=" * 50)
    report_lines.append("VAPT ANALYSIS REPORT")
    report_lines.append("=" * 50)
    report_lines.append("")
    
    report_lines.append("REPORT DETAILS:")
    report_lines.append(f"Auditor Name: {auditor_name}")
    report_lines.append(f"Certification: {audit_cert}")
    report_lines.append(f"Organization: {organization}")
    report_lines.append(f"Assessment Date: {assessment_date}")
    report_lines.append(f"Target OS: {selected_os}")
    report_lines.append(f"Overall Progress: {calculate_overall_progress()}%")
    report_lines.append("")
    
    if st.session_state.analysis_results:
        report_lines.append("ANALYSIS RESULTS:")
        report_lines.append("-" * 30)
        
        for section, result in st.session_state.analysis_results.items():
            if isinstance(result, dict):
                report_lines.append(f"\n{section}:")
                report_lines.append(f"Risk Level: {result.get('risk', 'N/A')}")
                report_lines.append(f"CVSS Score: {result.get('cvss', 'N/A')}")
                report_lines.append(f"Vulnerabilities: {result.get('vulnerabilities', 'N/A')}")
                report_lines.append(f"CVE References: {result.get('cve', 'N/A')}")
                report_lines.append(f"Framework References: {result.get('references', 'N/A')}")
                report_lines.append(f"Remediation: {result.get('remediation', 'N/A')}")
    
    return "\n".join(report_lines)

# Main Application
st.title("🔒 Your System VAPT Analysis")

# Header and Configuration Section
st.subheader("🔑 AI Analysis Configuration")
st.caption("💡 Add API key to enable AI analysis")

col1, col2 = st.columns([3, 1])
with col1:
    api_key = st.text_input("API Key (OpenAI, Gemini, Claude, Grok)", 
                           value=st.session_state.api_key, 
                           type="password",
                           help="Enter your AI provider API key for advanced analysis")
    if api_key != st.session_state.api_key:
        st.session_state.api_key = api_key

with col2:
    if st.button("Reset Assessment", type="secondary"):
        # Clear session state
        for key in list(st.session_state.keys()):
            if key != 'api_key':  # Keep API key
                del st.session_state[key]
        
        # Clear data folder
        if os.path.exists("data"):
            shutil.rmtree("data")
        os.makedirs("data", exist_ok=True)
        
        # Clear cache
        st.cache_data.clear()
        st.success("Assessment reset successfully!")
        st.rerun()

# Progress Dashboard
st.subheader("📊 Assessment Progress")
st.caption("Track your completion status across all vulnerability assessment categories")

progress_percentage = calculate_overall_progress()
st.progress(progress_percentage / 100)
st.write(f"Overall Progress: {progress_percentage}%")

# Create progress indicators
col1, col2 = st.columns(2)

with col1:
    st.write("**Windows 11 Categories:**")
    for section in WINDOWS_COMMANDS.keys():
        status = "✅" if st.session_state.progress.get(f"windows_{section}", False) else "⏳"
        st.write(f"{status} {section}")

with col2:
    st.write("**macOS Categories:**")
    for section in MACOS_COMMANDS.keys():
        status = "✅" if st.session_state.progress.get(f"macos_{section}", False) else "⏳"
        st.write(f"{status} {section}")

st.divider()

# OS Selection
st.subheader("🎯 Target System")

tab1, tab2 = st.tabs(["Windows 11", "macOS"])

def display_os_sections(os_name, commands_dict):
    """Display sections for the selected OS"""
    
    for section_name, commands in commands_dict.items():
        with st.expander(f"{section_name}", expanded=False):
            
            # Display commands table
            st.write("**Available Commands:**")
            
            # Create a more readable table
            for i, cmd in enumerate(commands, 1):
                st.write(f"**Command {i}: `{cmd['command']}`**")
                st.write(f"**Purpose:** {cmd['purpose']}")
                st.write(f"**Instructions:** {cmd['instructions']}")
                st.write(f"**Shortcuts:** {cmd['shortcuts']}")
                st.write("---")
            
            # Text area for command outputs
            section_key = f"{os_name.lower()}_{section_name}"
            output_key = f"output_{section_key}"
            
            if output_key not in st.session_state.command_outputs:
                st.session_state.command_outputs[output_key] = ""
            
            command_output = st.text_area(
                f"Paste command outputs for {section_name}:",
                value=st.session_state.command_outputs[output_key],
                height=200,
                key=f"textarea_{section_key}",
                help="Run the commands above on your system and paste the complete output here"
            )
            
            if command_output != st.session_state.command_outputs[output_key]:
                st.session_state.command_outputs[output_key] = command_output
            
            # Analysis buttons
            if command_output.strip():
                col1, col2 = st.columns(2)
                
                with col1:
                    if st.button(f"🧪 Test Run", key=f"mock_{section_key}"):
                        with st.spinner("Running mock analysis..."):
                            mock_result = mock_analyze_output(command_output, section_name)
                            st.session_state.analysis_results[section_key] = mock_result
                            st.session_state.progress[section_key] = True
                        st.success("Mock analysis completed!")
                
                with col2:
                    ai_disabled = not st.session_state.api_key.strip()
                    if st.button(f"🤖 AI Run", key=f"ai_{section_key}", disabled=ai_disabled):
                        if st.session_state.api_key.strip():
                            with st.spinner("Running AI analysis..."):
                                ai_result = analyze_with_ai(command_output, st.session_state.api_key, section_name)
                                st.session_state.analysis_results[section_key] = ai_result
                                st.session_state.progress[section_key] = True
                            st.success("AI analysis completed!")
                        else:
                            st.error("Please provide an API key for AI analysis")
                
                # Display analysis results
                if section_key in st.session_state.analysis_results:
                    result = st.session_state.analysis_results[section_key]
                    
                    st.subheader("🔍 Analysis Results")
                    
                    if isinstance(result, dict):
                        # Mock analysis result (structured)
                        col1, col2, col3 = st.columns(3)
                        
                        with col1:
                            risk_color = {"Low": "🟢", "Medium": "🟡", "High": "🔴"}.get(result.get('risk'), '⚪')
                            st.metric("Risk Level", f"{risk_color} {result.get('risk', 'N/A')}")
                        
                        with col2:
                            st.metric("CVSS Score", result.get('cvss', 'N/A'))
                        
                        with col3:
                            st.metric("CVE References", result.get('cve', 'N/A'))
                        
                        st.write(f"**Vulnerabilities Found:** {result.get('vulnerabilities', 'N/A')}")
                        st.write(f"**Framework References:** {result.get('references', 'N/A')}")
                        st.write(f"**Recommended Actions:** {result.get('remediation', 'N/A')}")
                    
                    else:
                        # AI analysis result (text)
                        st.write(result)
            
            else:
                st.info("Please paste command outputs above to enable analysis.")
    
    # Full Analysis Buttons Section
    st.divider()
    st.subheader("🚀 Bulk Analysis Options")
    st.info("💡 Run analysis on all sections that have command outputs at once")
    
    # Count sections with outputs
    sections_with_output = []
    for section_name, commands in commands_dict.items():
        section_key = f"{os_name.lower()}_{section_name}"
        output_key = f"output_{section_key}"
        if output_key in st.session_state.command_outputs and st.session_state.command_outputs[output_key].strip():
            sections_with_output.append((section_name, section_key, output_key))
    
    if sections_with_output:
        st.write(f"**Found {len(sections_with_output)} sections with command outputs ready for analysis:**")
        for section_name, _, _ in sections_with_output:
            st.write(f"✅ {section_name}")
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("🧪 Full Test Analysis", key=f"full_mock_{os_name.lower()}"):
                with st.spinner(f"Running test analysis on {len(sections_with_output)} sections..."):
                    success_count = 0
                    for section_name, section_key, output_key in sections_with_output:
                        try:
                            command_output = st.session_state.command_outputs[output_key]
                            mock_result = mock_analyze_output(command_output, section_name)
                            st.session_state.analysis_results[section_key] = mock_result
                            st.session_state.progress[section_key] = True
                            success_count += 1
                        except Exception as e:
                            st.error(f"Error analyzing {section_name}: {str(e)}")
                    
                st.success(f"✅ Test analysis completed for {success_count}/{len(sections_with_output)} sections!")
                st.balloons()
        
        with col2:
            ai_disabled = not st.session_state.api_key.strip()
            if st.button("🤖 Full AI Analysis", key=f"full_ai_{os_name.lower()}", disabled=ai_disabled):
                if st.session_state.api_key.strip():
                    with st.spinner(f"Running AI analysis on {len(sections_with_output)} sections..."):
                        success_count = 0
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        
                        for i, (section_name, section_key, output_key) in enumerate(sections_with_output):
                            try:
                                status_text.text(f"Analyzing {section_name}...")
                                command_output = st.session_state.command_outputs[output_key]
                                ai_result = analyze_with_ai(command_output, st.session_state.api_key, section_name)
                                st.session_state.analysis_results[section_key] = ai_result
                                st.session_state.progress[section_key] = True
                                success_count += 1
                                progress_bar.progress((i + 1) / len(sections_with_output))
                            except Exception as e:
                                st.error(f"Error analyzing {section_name}: {str(e)}")
                        
                        status_text.text("Analysis complete!")
                    st.success(f"✅ AI analysis completed for {success_count}/{len(sections_with_output)} sections!")
                    st.balloons()
                else:
                    st.error("Please provide an API key for AI analysis")
            
            if ai_disabled:
                st.caption("⚠️ API key required for AI analysis")
    else:
        st.warning("⚠️ No sections with command outputs found. Please paste command outputs in the sections above first.")

# Display content for selected tab
with tab1:
    st.write("### Windows 11 Vulnerability Assessment")
    st.info("💡 **Tip:** Run Command Prompt or PowerShell as Administrator for most commands")
    display_os_sections("Windows", WINDOWS_COMMANDS)

with tab2:
    st.write("### macOS Vulnerability Assessment") 
    st.info("💡 **Tip:** Some commands may require admin password (sudo)")
    display_os_sections("macOS", MACOS_COMMANDS)

st.divider()

# Report Generation Section
st.subheader("📋 Generate Assessment Report")

col1, col2 = st.columns(2)

with col1:
    auditor_name = st.text_input("Enter Auditor Name", placeholder="John Doe")
    audit_cert = st.text_input("Certificate #", placeholder="CISA, CISSP, CEH, etc.")

with col2:
    organization = st.text_input("Organisation", placeholder="Company/Organisation name")
    assessment_date = st.date_input("Assessment Date")

# Get selected OS from tabs (simplified - using first tab with data)
selected_os = "Windows 11"  # Default
if any(key.startswith('macos_') for key in st.session_state.progress.keys()):
    selected_os = "macOS"

# Report buttons
col1, col2, col3 = st.columns(3)

with col1:
    if st.button("👀 Preview Report"):
        if auditor_name and organization:
            st.subheader("📄 Report Preview")
            
            st.write("**Report Details:**")
            st.write(f"- **Auditor:** {auditor_name}")
            st.write(f"- **Certificate:** {audit_cert}")
            st.write(f"- **Organization:** {organization}")
            st.write(f"- **Date:** {assessment_date}")
            st.write(f"- **Target OS:** {selected_os}")
            st.write(f"- **Progress:** {calculate_overall_progress()}%")
            
            if st.session_state.analysis_results:
                st.write("**Analysis Summary:**")
                for section, result in st.session_state.analysis_results.items():
                    if isinstance(result, dict):
                        risk_icon = {"Low": "🟢", "Medium": "🟡", "High": "🔴"}.get(result.get('risk'), '⚪')
                        st.write(f"- **{section}:** {risk_icon} {result.get('risk', 'N/A')} Risk (CVSS: {result.get('cvss', 'N/A')})")
        else:
            st.error("Please fill in auditor name and organization")

with col2:
    if st.button("📄 Generate PDF"):
        if auditor_name and organization:
            with st.spinner("Generating PDF report..."):
                pdf_buffer = generate_pdf_report(auditor_name, audit_cert, organization, 
                                                str(assessment_date), selected_os)
                
                st.download_button(
                    label="⬇️ Download PDF Report",
                    data=pdf_buffer,
                    file_name=f"VAPT_Report_{organization}_{assessment_date}.pdf",
                    mime="application/pdf"
                )
        else:
            st.error("Please fill in auditor name and organization")

with col3:
    if st.button("📝 Generate Word"):
        if auditor_name and organization:
            with st.spinner("Generating Word document..."):
                docx_buffer = generate_docx_report(auditor_name, audit_cert, organization,
                                                  str(assessment_date), selected_os)
                
                st.download_button(
                    label="⬇️ Download Word Report",
                    data=docx_buffer,
                    file_name=f"VAPT_Report_{organization}_{assessment_date}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.error("Please fill in auditor name and organization")

# Text report
if st.button("📰 Generate Text Report"):
    if auditor_name and organization:
        text_report = generate_text_report(auditor_name, audit_cert, organization,
                                         str(assessment_date), selected_os)
        
        st.download_button(
            label="⬇️ Download Text Report",
            data=text_report,
            file_name=f"VAPT_Report_{organization}_{assessment_date}.txt",
            mime="text/plain"
        )
    else:
        st.error("Please fill in auditor name and organization")

# Footer
st.divider()
st.caption("🔒 **Security Notice:** This tool is for authorized security assessments only. Ensure you have proper knowledge before running these commands on any system.")
st.caption("⚠️ **Disclaimer:** This application is developed as an ICAI AICA L2 Capstone Project for educational purposes only. Neither ICAI, its AI Committee, nor the creator shall be responsible for any misuse or damage arising from its use. Please exercise caution and validate results independently")
